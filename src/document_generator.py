from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

class StrategyDocumentGenerator:
    """Generate professional DOCX documents for content strategies"""

    def __init__(self):
        self.doc = Document()
        self.setup_styles()

    def setup_styles(self):
        """Configure document styles"""
        # Set default font
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

    def add_title_page(self, brand_name, strategy_count=5):
        """Add professional title page"""
        # Title
        title = self.doc.add_heading('AI Content Marketing Strategy', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f'{strategy_count} Strategic Approaches for {brand_name}')
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(74, 144, 226)  # Blue

        # Date
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
        date_run.font.size = Pt(12)
        date_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray

        # Add some space
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Divider
        self.doc.add_paragraph('_' * 60)

        # Page break
        self.doc.add_page_break()

    def add_section(self, title, content, level=1):
        """Add a section with heading and content"""
        self.doc.add_heading(title, level=level)

        if isinstance(content, str):
            self.doc.add_paragraph(content)
        elif isinstance(content, list):
            for item in content:
                self.doc.add_paragraph(item, style='List Bullet')
        elif isinstance(content, dict):
            for key, value in content.items():
                p = self.doc.add_paragraph()
                p.add_run(f'{key}: ').bold = True
                p.add_run(str(value))

    def add_strategy(self, strategy_num, strategy_data):
        """Add a complete strategy section"""
        # Strategy header
        self.doc.add_heading(f'Strategy {strategy_num}: {strategy_data.get("name", "Unnamed Strategy")}', level=1)

        # Tagline
        if 'tagline' in strategy_data:
            tagline = self.doc.add_paragraph()
            tagline_run = tagline.add_run(f'"{strategy_data["tagline"]}"')
            tagline_run.italic = True
            tagline_run.font.size = Pt(12)
            tagline_run.font.color.rgb = RGBColor(74, 144, 226)

        self.doc.add_paragraph()  # Space

        # Core approach
        if 'core_approach' in strategy_data:
            self.doc.add_heading('Core Approach', level=2)
            self.doc.add_paragraph(strategy_data['core_approach'])

        # Content Pillars
        if 'content_pillars' in strategy_data:
            self.doc.add_heading('Content Pillars', level=2)
            pillars = strategy_data['content_pillars']
            if isinstance(pillars, list):
                for pillar in pillars:
                    if isinstance(pillar, dict):
                        p = self.doc.add_paragraph(style='List Bullet')
                        p.add_run(f"{pillar.get('name', 'Unnamed Pillar')}: ").bold = True
                        p.add_run(pillar.get('description', ''))
                    else:
                        self.doc.add_paragraph(str(pillar), style='List Bullet')

        # Posting Frequency
        if 'posting_frequency' in strategy_data:
            self.doc.add_heading('Posting Frequency', level=2)
            freq = strategy_data['posting_frequency']
            if isinstance(freq, dict):
                for channel, count in freq.items():
                    self.doc.add_paragraph(f'{channel}: {count}', style='List Bullet')
            else:
                self.doc.add_paragraph(str(freq))

        # Content Mix
        if 'content_mix' in strategy_data:
            self.doc.add_heading('Content Mix', level=2)
            mix = strategy_data['content_mix']
            if isinstance(mix, dict):
                for content_type, percentage in mix.items():
                    self.doc.add_paragraph(f'{content_type}: {percentage}%', style='List Bullet')

        # Top Content Ideas
        if 'top_5_ideas' in strategy_data:
            self.doc.add_heading('Top 5 Content Ideas', level=2)
            for i, idea in enumerate(strategy_data['top_5_ideas'], 1):
                self.doc.add_paragraph(f'{i}. {idea}', style='List Number')

        # Expected Results
        if 'expected_results' in strategy_data:
            self.doc.add_heading('Expected 30-Day Results', level=2)
            results = strategy_data['expected_results']
            if isinstance(results, list):
                for result in results:
                    self.doc.add_paragraph(result, style='List Bullet')
            else:
                self.doc.add_paragraph(str(results))

        # Pros and Cons
        self.doc.add_heading('Pros & Cons', level=2)

        if 'pros' in strategy_data:
            self.doc.add_paragraph().add_run('Pros:').bold = True
            pros = strategy_data['pros']
            if isinstance(pros, list):
                for pro in pros:
                    p = self.doc.add_paragraph(pro, style='List Bullet')
                    p.runs[0].font.color.rgb = RGBColor(0, 128, 0)  # Green

        if 'cons' in strategy_data:
            self.doc.add_paragraph().add_run('Cons:').bold = True
            cons = strategy_data['cons']
            if isinstance(cons, list):
                for con in cons:
                    p = self.doc.add_paragraph(con, style='List Bullet')
                    p.runs[0].font.color.rgb = RGBColor(255, 69, 0)  # Red

        # Page break after each strategy
        self.doc.add_page_break()

    def add_recommendation(self, recommendation_text):
        """Add recommendation section"""
        self.doc.add_heading('Our Recommendation', level=1)

        # Highlight box effect with paragraph
        rec_para = self.doc.add_paragraph()
        rec_run = rec_para.add_run('💡 ' + recommendation_text)
        rec_run.font.size = Pt(12)

        self.doc.add_paragraph()  # Space

    def save(self, filepath):
        """Save the document"""
        self.doc.save(filepath)
        return filepath


class CalendarDocumentGenerator:
    """Generate professional DOCX documents for content calendars"""

    def __init__(self):
        self.doc = Document()
        self.setup_styles()

    def setup_styles(self):
        """Configure document styles"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

    def add_title_page(self, brand_name, strategy_name, month):
        """Add title page for calendar"""
        title = self.doc.add_heading('Content Calendar', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f'{brand_name} • {strategy_name}')
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(74, 144, 226)

        month_para = self.doc.add_paragraph()
        month_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        month_run = month_para.add_run(month)
        month_run.font.size = Pt(14)

        self.doc.add_paragraph('_' * 60)
        self.doc.add_page_break()

    def add_executive_summary(self, summary):
        """Add executive summary"""
        self.doc.add_heading('Executive Summary', level=1)
        self.doc.add_paragraph(summary)
        self.doc.add_paragraph()

    def add_content_piece(self, piece_data, brand_name=''):
        """Add a single content piece"""
        # Content header
        header = self.doc.add_heading(
            f"Content #{piece_data.get('content_id', '?')}: {piece_data.get('title', 'Untitled')}",
            level=2
        )

        # Details table-like format
        details = [
            ('Week', piece_data.get('week', 'TBD')),
            ('Date', piece_data.get('suggested_date', 'TBD')),
            ('Channel', piece_data.get('channel', 'TBD')),
            ('Format', piece_data.get('format', 'TBD')),
            ('Pillar', piece_data.get('pillar', 'TBD')),
        ]

        for label, value in details:
            p = self.doc.add_paragraph()
            p.add_run(f'{label}: ').bold = True
            p.add_run(str(value))

        self.doc.add_paragraph()  # Space

        # Description - ENSURE IT'S NOT EMPTY
        description = piece_data.get('description', '')
        if not description or description in ['Content description', 'Engaging content', 'Description']:
            title = piece_data.get('title', 'this topic')
            format_type = piece_data.get('format', 'engaging content')
            description = f"Detailed content about {title} for {brand_name}. This piece will provide value to the audience through {format_type}."

        desc_para = self.doc.add_paragraph()
        desc_para.add_run('Description: ').bold = True
        desc_para.add_run(description)

        # Key Message - ENSURE IT'S NOT EMPTY
        key_message = piece_data.get('key_message', '')
        if not key_message or key_message in ['Key message', 'Main message']:
            pillar = piece_data.get('pillar', 'brand values')
            key_message = f"Main message highlighting {brand_name}'s value proposition related to {pillar}."

        msg_para = self.doc.add_paragraph()
        msg_para.add_run('Key Message: ').bold = True
        msg_para.add_run(key_message)

        # Call to Action
        cta = piece_data.get('call_to_action', 'Take action')
        cta_para = self.doc.add_paragraph()
        cta_para.add_run('Call to Action: ').bold = True
        cta_para.add_run(cta)

        # Effort and Engagement
        meta = self.doc.add_paragraph()
        meta.add_run('Effort: ').bold = True
        meta.add_run(f"{piece_data.get('effort_level', 'Medium')} • ")
        meta.add_run('Engagement Potential: ').bold = True
        meta.add_run(piece_data.get('engagement_potential', 'Medium'))

        # Execution Notes - ENSURE IT'S NOT EMPTY
        exec_notes = piece_data.get('execution_notes', '')
        if not exec_notes or exec_notes in ['Execution notes', 'Tips for creating content piece', 'Notes']:
            format_type = piece_data.get('format', 'content')
            channel = piece_data.get('channel', 'the platform')
            effort = piece_data.get('effort_level', 'medium')
            exec_notes = f"Production tips: Plan {format_type} creation for {channel}. Consider {effort} effort level when scheduling production."

        notes_para = self.doc.add_paragraph()
        notes_para.add_run('Execution Notes: ').bold = True
        notes_run = notes_para.add_run(exec_notes)
        notes_run.font.italic = True
        notes_run.font.color.rgb = RGBColor(100, 100, 100)

        self.doc.add_paragraph()  # Space between pieces

    def save(self, filepath):
        """Save the document"""
        self.doc.save(filepath)
        return filepath


def generate_strategy_docx(brand_name, strategies_list, recommendation, output_path):
    """
    Generate a DOCX document with all 5 strategies

    Args:
        brand_name: Name of the brand
        strategies_list: List of strategy dictionaries
        recommendation: Recommendation text
        output_path: Where to save the file
    """
    gen = StrategyDocumentGenerator()
    gen.add_title_page(brand_name, len(strategies_list))

    for i, strategy in enumerate(strategies_list, 1):
        gen.add_strategy(i, strategy)

    gen.add_recommendation(recommendation)

    return gen.save(output_path)


def generate_calendar_docx(brand_name, strategy_name, month, executive_summary, content_pieces, output_path):
    """
    Generate a DOCX document for content calendar

    Args:
        brand_name: Name of the brand
        strategy_name: Selected strategy name
        month: Month (e.g., "January 2025")
        executive_summary: Summary text
        content_pieces: List of content piece dictionaries
        output_path: Where to save the file
    """
    gen = CalendarDocumentGenerator()
    gen.add_title_page(brand_name, strategy_name, month)
    gen.add_executive_summary(executive_summary)

    for piece in content_pieces:
        gen.add_content_piece(piece, brand_name=brand_name)

    return gen.save(output_path)
